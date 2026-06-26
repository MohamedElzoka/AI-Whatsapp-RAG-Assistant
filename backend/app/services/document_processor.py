"""
Knowledge-base document pipeline: text extraction + chunking.

Supported formats: PDF, DOCX, TXT (per design doc).
"""
import re

import docx
from pypdf import PdfReader

from app.config import settings


class UnsupportedFileTypeError(Exception):
    pass


def extract_text(file_path: str, file_type: str) -> str:
    file_type = file_type.lower().lstrip(".")
    if file_type == "pdf":
        return _extract_pdf(file_path)
    if file_type == "docx":
        return _extract_docx(file_path)
    if file_type == "txt":
        return _extract_txt(file_path)
    raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]
    # Also pull text out of any tables (common in policy / pricing docs)
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def clean_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def chunk_text(
    text: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[str]:
    """
    Simple, dependency-free recursive-ish splitter: splits on paragraph
    boundaries first, then falls back to sliding-window character chunks
    so we never produce an oversized chunk that blows the embedding limit.
    """
    chunk_size = chunk_size or settings.CHUNK_SIZE
    chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP

    text = clean_text(text)
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    chunks: list[str] = []
    buffer = ""

    for para in paragraphs:
        candidate = f"{buffer}\n\n{para}" if buffer else para

        if len(candidate) <= chunk_size:
            buffer = candidate
            continue

        if buffer:
            chunks.append(buffer)

        if len(para) <= chunk_size:
            buffer = para
        else:
            # Paragraph itself is too long: slide a window across it
            for start in range(0, len(para), chunk_size - chunk_overlap):
                chunks.append(para[start : start + chunk_size])
            buffer = ""

    if buffer:
        chunks.append(buffer)

    return chunks
